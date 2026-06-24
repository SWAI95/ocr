"""OCR 엔진/모델 비교 보고서(.docx) 생성기.

doc_00(한국어 계약서 5페이지) 실측 결과 + 엔진/모델 메타데이터를 모아
초보·준전문가가 함께 볼 수 있는 Word 보고서를 만든다.

실행: ./.venv/bin/python scripts/make_report.py
산출: OCR_엔진_모델_비교_보고서.docx (프로젝트 루트)
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "OCR_엔진_모델_비교_보고서.docx"

# 측정 데이터(SSOT) — scripts/measure_all.py 산출물. CPU·도장 표를 여기서 채운다.
_MEAS = Path(__file__).resolve().parent.parent / "data" / "measurements.json"
M = json.loads(_MEAS.read_text(encoding="utf-8")) if _MEAS.exists() else {}


def meas_cell(v, suf="%"):
    return f"{v:.1f}{suf}" if isinstance(v, (int, float)) else "측정중"

# ── 색상 ────────────────────────────────────────────────────────────────
NAVY = RGBColor(0x1F, 0x33, 0x55)
GREEN = RGBColor(0x1B, 0x7A, 0x3D)
RED = RGBColor(0xB4, 0x23, 0x23)
GRAY = RGBColor(0x60, 0x60, 0x60)
HDR_BG = "1F3355"   # 표 헤더 배경(네이비)
BEST_BG = "E7F3EA"  # 베스트 행 배경(연녹)


# ── 문서 기본 폰트(한글) ──────────────────────────────────────────────
def _set_base_font(doc: Document) -> None:
    st = doc.styles["Normal"]
    st.font.name = "맑은 고딕"
    st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def _shade(cell, hex_color: str) -> None:
    tc = cell._tc.get_or_add_tcPr()
    shd = tc.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color})
    tc.append(shd)


def _set_cell(cell, text, *, bold=False, color=None, white=False, size=9, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color:
        run.font.color.rgb = color


def _header_row(table, headers):
    for j, h in enumerate(headers):
        c = table.rows[0].cells[j]
        _shade(c, HDR_BG)
        _set_cell(c, h, bold=True, white=True, size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER)


def add_table(doc, headers, rows, *, widths=None, best_row=None, center_cols=()):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _header_row(t, headers)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            al = WD_ALIGN_PARAGRAPH.CENTER if j in center_cols else None
            _set_cell(cells[j], val, size=9, align=al,
                      bold=(best_row is not None and ri == best_row))
            if best_row is not None and ri == best_row:
                _shade(cells[j], BEST_BG)
    if widths:
        for j, w in enumerate(widths):
            for r in t.rows:
                r.cells[j].width = w
    return t


def h(doc, text, level):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = "맑은 고딕"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        if level <= 1:
            r.font.color.rgb = NAVY
    return p


def para(doc, text, *, bold=False, color=None, size=10, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def bullets(doc, items, *, style="List Bullet"):
    for it in items:
        p = doc.add_paragraph(style=style)
        # "라벨: 내용" 이면 라벨만 볼드
        if isinstance(it, tuple):
            label, rest = it
            r1 = p.add_run(label)
            r1.bold = True
            p.add_run(rest)
        else:
            p.add_run(it)


# ════════════════════════════════════════════════════════════════════════
doc = Document()
_set_base_font(doc)

# ── 표지 ────────────────────────────────────────────────────────────────
title = doc.add_heading("오프라인 한국어 문서 OCR\n엔진·모델 비교 보고서", level=0)
for r in title.runs:
    r.font.name = "맑은 고딕"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    r.font.color.rgb = NAVY
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = sub.add_run("테스트셋: 한국어 계약서 doc_00·doc_04·doc_05 (총 28페이지)  ·  지표: CER/WER  ·  평가 기준 GT 대비")
rr.font.size = Pt(10)
rr.font.color.rgb = GRAY

para(doc, "")
box = doc.add_paragraph()
box.alignment = WD_ALIGN_PARAGRAPH.CENTER
b1 = box.add_run("결론 한 줄:  ")
b1.bold = True
b1.font.size = Pt(12)
b2 = box.add_run("하이브리드(hybrid_vl) 엔진이 3개 문서 모두 CER 최저(평균 1.26%)로 종합 1위 — 한국어 문서 OCR 베스트")
b2.font.size = Pt(12)
b2.font.color.rgb = GREEN
b2.bold = True

# ── 0. 개요 ─────────────────────────────────────────────────────────────
h(doc, "0. 이 문서의 목적과 읽는 법", 1)
para(doc, "본 보고서는 한 종류의 OCR 모델이 아니라 여러 모델을 조합한 \"엔진\" 6종을 "
          "사용자 입장에서 비교한다. 사용자는 결국 \"어떤 엔진을 고를까\"를 결정하므로, "
          "① 엔진 비교 표로 큰 그림을 잡고 ② 성능 표로 수치를 확인한 뒤 ③ 엔진별·모델별 "
          "상세 장표에서 근거를 본다. 마지막에 베스트 엔진을 정리한다.")
para(doc, "초보자를 위한 용어 (아래 표를 읽기 전에 확인):", bold=True)
add_table(
    doc,
    ["용어", "쉽게 말하면", "좋고 나쁨"],
    [
        ["CER", "글자 단위 오류율. 100자 중 몇 글자가 틀렸나(공백 무시)", "낮을수록 좋음 (1차 지표)"],
        ["WER", "띄어쓰기(단어) 단위 오류율. 가독성·문장 복원력", "낮을수록 좋음 (한글은 참고용)"],
        ["VLM", "이미지를 통째로 \"보고 읽는\" AI(비전언어모델). 예: Qwen2.5-VL", "줄바꿈·띄어쓰기 자연스러움"],
        ["결정론", "같은 입력 → 항상 같은 출력(예측 가능)", "✓ 사고 위험 없음 / ✗ 가끔 다른 결과"],
        ["오프라인", "인터넷 없는 내부망에서 동작 가능 여부", "✓ 내부망 배포 가능"],
    ],
    center_cols=(0,),
)
para(doc, "※ 한국어는 띄어쓰기 표기가 사람마다 달라 WER이 흔들린다. 그래서 글자 정확도인 "
          "CER을 1차 지표로 보고, WER(가독성)은 보조로 본다.", italic=True, color=GRAY, size=9)

# ── 1. 엔진 비교 표 (메인) ───────────────────────────────────────────────
h(doc, "1. 엔진 비교 표 (사용자 선택 기준)", 1)
para(doc, "각 엔진은 여러 모델의 조합이다. \"구성(모델 조합)\"이 그 엔진의 정체성이다.")
add_table(
    doc,
    ["엔진", "구성 (모델 조합)", "한 줄 특징", "추천도"],
    [
        ["hybrid_vl ★", "PP-DocLayout 분류 → 폼·표=PP-StructureV3 / 줄글=Qwen2.5-VL + 누락감지 병합",
         "폼·표는 Paddle, 줄글은 VLM. 사고 페이지 자동 구제. CER·WER 모두 최고", "◎ 1순위"],
        ["paddle", "PP-StructureV3 (PP-DocLayout + PP-OCRv5 + SLANeXt 표)",
         "결정론적·폼/표 구조 최강. 단 줄글 띄어쓰기 약함(WER↑)", "○ 표·폼 위주"],
        ["ollama_vl", "Qwen2.5-VL (로컬 Ollama)",
         "줄글 가독성 좋고 5090·4090 모두 GPU로 빠름. 박스 없음·가끔 사고", "○ 빠른 줄글"],
        ["paddle_vl", "PaddleOCR-VL-0.9B + PP-DocLayoutV3",
         "단독 VLM(서버 불필요), 5090·4090 GPU. reading-order 강함. hybrid의 오프라인 VLM 백엔드", "○ 오프라인 VLM"],
        ["easyocr", "CRAFT(검출) + CRNN(korean_g2)",
         "가볍고 빠르며 설치 쉬움. 문서 정확도는 낮음", "△ 간단/대량"],
        ["tesseract", "LSTM (kor+eng)",
         "초경량 오픈소스 baseline. 한국어 문서 정확도 낮음", "△ 대조군"],
    ],
    best_row=0,
    center_cols=(3,),
)

# ── 2. 성능 비교 표 (별도) ───────────────────────────────────────────────
h(doc, "2. 성능 비교 (doc_00·04·05 실측)", 1)
para(doc, "정확도(문서별 CER/WER)와 엔진 특성(속도·GPU·오프라인)을 나눠 표기. CER 낮을수록 좋음.")
para(doc, "2-A. 문서별 정확도 (CER / WER, %)", bold=True, size=10)


def _cw(e, s):  # "cer / wer" 셀 (측정 JSON)
    v = M.get("gpu", {}).get(s, {}).get(e)
    return f"{v['cer']:.1f} / {v['wer']:.1f}" if v else "—"


def _avg(e):
    cs = [M["gpu"][s][e]["cer"] for s in ("doc_00", "doc_04", "doc_05")
          if M.get("gpu", {}).get(s, {}).get(e)]
    return f"{sum(cs)/len(cs):.2f}%" if cs else "—"


_2a_rows = [[e + (" ★" if e == "hybrid_vl" else ""),
            _cw(e, "doc_00"), _cw(e, "doc_04"), _cw(e, "doc_05"), _avg(e)]
           for e in ["hybrid_vl", "paddle", "ollama_vl", "paddle_vl", "easyocr", "tesseract"]]
add_table(doc, ["엔진", "doc_00", "doc_04", "doc_05", "평균 CER"], _2a_rows,
          best_row=0, center_cols=(1, 2, 3, 4))
para(doc, "doc_00=계약서 5p · doc_04=도급계약서 11p · doc_05=계약서 12p (총 28p). "
          "hybrid_vl이 3개 문서 모두 CER 최저(평균 1.26%) → 종합 1위. paddle_vl(5090 GPU 실측)은 "
          "VLM이라 WER이 paddle보다 압도적으로 낮고(doc_00 16.7 vs 43.2) CER도 doc_04·05는 "
          "우수하나, doc_00은 상단 헤더(회사명·번호 로고) 누락으로 7.46%. hybrid는 그 헤더를 "
          "Paddle로 잡아 1.52%로 회복. paddle·paddle_vl·hybrid_vl은 이번 세션 재측정(GPU), "
          "ollama_vl·easyocr·tesseract는 직전 측정값.", italic=True, color=GRAY, size=9)
para(doc, "2-B. 엔진 특성 (문서 무관)", bold=True, size=10)
add_table(
    doc,
    ["엔진", "속도(5090)", "GPU 5090 / 4090", "오프라인", "결정론"],
    [
        ["hybrid_vl ★", "7~20초/p", "전체 / 전체", "✓(Ollama)", "~(seed)"],
        ["paddle", "0.4~1초/p", "GPU / GPU", "✓", "✓"],
        ["ollama_vl", "6~8초/p", "GPU / GPU", "✓(Ollama)", "~(seed)"],
        ["easyocr", "2~4초/p", "GPU / GPU", "✓", "✓"],
        ["paddle_vl", "5~20초/p", "GPU / GPU", "✓ 가중치", "✓(greedy)"],
        ["tesseract", "1.7초/p", "CPU / CPU", "✓", "✓"],
    ],
    center_cols=(1, 2, 3, 4),
)
para(doc, "※ 수치는 문서 단위(micro: 각 문서 전체 페이지를 이어붙여 한 번 계산)로 웹 멀티엔진 비교와 동일 방식. "
          "paddle·easyocr·tesseract는 결정론 → 웹과 정확히 일치. ollama_vl·hybrid_vl은 Ollama seed "
          "고정(OCR_OLLAMA_SEED, 기본 0)으로 변동을 크게 줄였으나 Ollama 잔여 비결정성으로 호출마다 "
          "±0.1%p 차이 가능('~' 표기) — 웹 값과 소수점 차이는 정상.",
     italic=True, color=GRAY, size=9)
para(doc, "※ 측정 환경: 개발기 RTX 5090(Blackwell) + paddlepaddle-gpu cu129(CUDA 12.9). cu129 "
          "휠로 Paddle도 5090에서 GPU 동작(0.4초/p) — 예전 cu126 휠은 Blackwell 미지원으로 CPU "
          "폴백했음. paddle_vl도 cu129 + markdown 접근자 수정 후 5090 GPU 정상 동작(과거 '빈 출력'은 "
          "Blackwell 한계가 아니라 결과 접근 버그였음). tesseract는 바이너리 설치 후 측정.",
     italic=True, color=GRAY, size=9)
para(doc, "※ 후처리 정규화(전 엔진 일괄): 숫자 앞 단독 'W'를 원화기호 '₩'로 교정(Paddle이 "
          "₩의 가로줄을 못 읽어 'W'로 인식 — KRW·MW 등은 보호) + VLM의 리터럴 '\\n' 아티팩트 제거. "
          "OCR_NORMALIZE=off로 비활성. doc_00 금액 페이지 CER이 이로써 소폭 개선됨.",
     italic=True, color=GRAY, size=9)
para(doc, "※ hybrid_vl 속도(~20초/p)의 대부분은 이제 Paddle이 아니라 Qwen(VLM) 호출 시간이다. "
          "안전장치의 Paddle 참조는 GPU로 ~0.4초라 더 이상 병목이 아니다.",
     italic=True, color=GRAY, size=9)
para(doc, "※ Ollama 엔진(ollama_vl·hybrid_vl)도 완전 오프라인이다 — qwen2.5vl 모델(~6GB)을 "
          "로컬에 받아둔 것으로 추론하며 인터넷 불필요(본 측정도 로컬 모델 사용). 내부망 배포 "
          "시엔 Ollama 데몬 설치 + 모델 반입(~6GB)만 추가하면 됨(파이썬 패키지로 자족하는 "
          "paddle/easyocr 과의 유일한 차이).",
     italic=True, color=GRAY, size=9)

# 2-C. GPU vs CPU (doc_00) — 측정 JSON 소비
para(doc, "2-C. GPU vs CPU 속도 (doc_00, 같은 문서·페이지당)", bold=True, size=10)
_cpu_rows = []
for _e in ["paddle", "easyocr", "tesseract", "ollama_vl", "paddle_vl", "hybrid_vl"]:
    _g = (M.get("gpu", {}).get("doc_00", {}).get(_e) or {})
    _c = (M.get("cpu_doc00", {}).get(_e) or {})
    _gs = f"{_g['per_page_s']}초" if _g.get("per_page_s") is not None else "—"
    _cs = f"{_c['per_page_s']}초" if _c.get("per_page_s") is not None else "측정중"
    _sp = (f"{_c['per_page_s']/max(_g['per_page_s'],0.01):.0f}배 느림"
           if _g.get("per_page_s") and _c.get("per_page_s") else "—")
    _cpu_rows.append([_e, _gs, _cs, _sp])
add_table(doc, ["엔진", "GPU 초/p", "CPU 초/p", "CPU 둔화"], _cpu_rows, center_cols=(1, 2, 3))
para(doc, "※ 딥러닝 엔진(paddle 등)은 CPU에서도 실용적이나, VLM(paddle_vl·ollama_vl·hybrid_vl)은 "
          "CPU에서 페이지당 수십~수백 초로 사실상 비실용 → 빠른 응답엔 GPU 사실상 필수.",
     italic=True, color=GRAY, size=9)

# 2-D. 도장 억제(red_stamp) on/off — 측정 JSON 소비
para(doc, "2-D. 도장 억제(빨강 제거) on/off CER (%, 낮을수록 좋음)", bold=True, size=10)
_seal_rows = []
for _stem, _lbl in [("doc_04", "doc_04(도장이 글자 가림)"), ("doc_00", "doc_00(도장 방해 안 함)")]:
    for _e in ["hybrid_vl", "paddle", "paddle_vl"]:
        _d = M.get("seal", {}).get(_stem, {}).get(_e, {})
        _off, _on = _d.get("off_cer"), _d.get("on_cer")
        _delta = (f"{_on-_off:+.1f}%p" if _off is not None and _on is not None else "측정중")
        _seal_rows.append([_lbl, _e, meas_cell(_off), meas_cell(_on), _delta])
add_table(doc, ["문서", "엔진", "억제 전", "억제 후", "변화"], _seal_rows, center_cols=(2, 3, 4))
para(doc, "※ 도장이 글자를 가리는 문서(doc_04)는 억제로 개선(−), 도장이 방해하지 않는 "
          "문서(doc_00)는 억제로 오히려 저하(+). → 전역 적용 금물, 선택적 적용해야 함.",
     italic=True, color=GRAY, size=9)

# ── 3. 엔진별 상세 장표 ──────────────────────────────────────────────────
h(doc, "3. 엔진별 상세 장표", 1)

ENGINES = [
    {
        "name": "hybrid_vl  (이 프로젝트 기본 엔진) ★",
        "combo": "PP-DocLayout(레이아웃 분류) → 폼·표 페이지=PP-StructureV3 / 줄글 페이지=Qwen2.5-VL, "
                 "+ 누락감지 병합 안전장치",
        "how": "페이지를 먼저 레이아웃으로 분류해, 표·폼이 많으면 Paddle(구조·글자 정확), 줄글이 "
               "많으면 Qwen2.5-VL(자연스러운 줄바꿈)로 보낸다. Qwen이 밀집 구간을 통째로 누락·"
               "전치하는 사고를 감지하면(hole≥35) 그 페이지만 \"Paddle 글자 + Qwen 띄어쓰기\"로 "
               "병합해 구제한다.",
        "pros": ["3개 문서 평균 CER 1.26%(전체 최저) · WER 4~15%(가독성 상위)로 종합 1위",
                 "폼/표는 구조 정확(Paddle), 줄글은 가독성(VLM) — 장점만 취함",
                 "줄글은 항상 'Paddle 글자+Qwen 띄어쓰기' 병합 → Qwen 사고에 강건, CER 안정(~1.7%, ±0.1)",
                 "VLM 백엔드 선택(OCR_HYBRID_VLM): auto(기본)/ollama_vl(Qwen)/paddle_vl. "
                 "paddle_vl 백엔드는 Ollama 없이 동작하는 완전 오프라인 경로(평균 CER 2.3%)"],
        "cons": ["dev(5090) ~20초/p: 대부분 Qwen(VLM) 호출 시간(Paddle 참조는 GPU로 빠름)",
                 "Ollama 서버 + Paddle 모델 둘 다 필요",
                 "줄글 페이지 일부는 여전히 VLM 의존"],
        "fit": "한국어 실문서 전반(계약서·공문 등 폼+표+줄글이 섞인 문서). 정확도 최우선일 때.",
    },
    {
        "name": "paddle  (Primary, PP-StructureV3)",
        "combo": "PP-DocLayout(레이아웃) + PP-OCRv5(검출/한국어 인식) + SLANeXt(표 구조)",
        "how": "레이아웃·텍스트·표·도장을 한 파이프라인에서 처리. 박스 좌표가 있어 오버레이·"
               "검수에 유리. 완전 오프라인.",
        "pros": ["결정론적 — 같은 입력에 항상 같은 결과(사고 위험 없음)",
                 "폼·표 구조 인식 최강(셀 정렬·도장 등), CER 2.28%로 우수",
                 "완전 오프라인(외부 서버 불필요), 박스 좌표 제공"],
        "cons": ["한국어 줄글 띄어쓰기를 뭉갬 → WER 43.2%(가독성 약함)",
                 "양쪽정렬 줄글에서 읽기순서가 흔들릴 수 있음",
                 "예전 cu126 휠은 5090서 CPU였으나 cu129 휠로 교체 시 5090도 GPU(0.4초/p)"],
        "fit": "표·폼 위주 문서, 결정론·오프라인이 필수인 환경, 박스 좌표가 필요한 검수용.",
    },
    {
        "name": "ollama_vl  (Qwen2.5-VL, 로컬)",
        "combo": "Qwen2.5-VL 7B (로컬 Ollama 서버)",
        "how": "문서 이미지를 통째로 VLM에 넣어 마크다운/표로 받아온다. Ollama의 자체 CUDA 빌드라 "
               "5090(Blackwell)에서도 GPU로 동작(Paddle의 sm_120 문제 우회).",
        "pros": ["줄바꿈·띄어쓰기가 자연스러움 → WER 13.6%(가독성 좋음)",
                 "5090·4090 모두 GPU로 빠름(8.4초/p)",
                 "Blackwell 우회 — dev에서도 GPU 사용"],
        "cons": ["밀집 구간 전치·누락·환각 가능(CER ~2.8%) — seed로 ±0.1 안정화되나 오류 자체는 남음",
                 "박스 좌표 없음(오버레이·정밀 검수 불리)",
                 "표 구조 정확도는 Paddle보다 약함, Ollama 서버 필요"],
        "fit": "줄글 위주 문서의 빠른 인식, GPU 가속이 필요한 dev/배포, 가독성 우선.",
    },
    {
        "name": "paddle_vl  (PaddleOCR-VL-0.9B)",
        "combo": "PaddleOCR-VL-1.6-0.9B(VLM) + PP-DocLayoutV3",
        "how": "0.9B 경량 VLM을 in-process(native)로 실행 — 별도 추론 서버 불필요. 이미지 → "
               "레이아웃 + 블록별 VLM 인식 → 마크다운/표. cu129 휠 + markdown 접근자 수정 후 "
               "5090(Blackwell)·4090 모두 GPU로 동작(4.8~11.5초/p).",
        "pros": ["단독 VLM으로 레이아웃+OCR(외부 서빙 백엔드 불필요), 완전 in-process",
                 "5090·4090 모두 GPU 네이티브(서버 불필요한 오프라인 VLM)",
                 "reading-order가 깨지는 폼/줄글에 강함(doc_05_p9 paddle 19.6→VL 3.3%)"],
        "cons": ["표-heavy·헤더 페이지엔 약함(표가 <table>로 빠져 라인 누락, 헤더 로고 미독)",
                 "단독 평균 CER은 paddle보다 약간 높음(3.9% vs 3.1%) → 라우팅 결합이 정답",
                 "VLM이라 박스 좌표 없음, 가중치 사전 다운로드 필요"],
        "fit": "hybrid_vl의 오프라인 VLM 백엔드(OCR_HYBRID_VLM=paddle_vl) — Ollama 없이 "
               "reading-order 복원. 5090/4090 단독 VLM 인식.",
    },
    {
        "name": "easyocr  (CRAFT + CRNN)",
        "combo": "CRAFT(텍스트 검출) + CRNN korean_g2(인식)",
        "how": "전통적 2단계 OCR(검출→인식). PyTorch 기반이라 5090에서도 GPU 동작. 레이아웃·표 "
               "구조 인식은 없음.",
        "pros": ["가볍고 빠름(3.5초/p), 설치 간단, 완전 오프라인",
                 "5090·4090 모두 GPU(torch cu128)",
                 "간단한 텍스트/짧은 문서에 무난"],
        "cons": ["조밀 단락·표 문서 정확도 낮음(CER 11.4%)",
                 "레이아웃/표 구조 없음, 줄바꿈 복원 약함(WER 45.0%)",
                 "한국어 복잡 문서엔 부적합"],
        "fit": "정확도보다 속도·간편함이 중요한 대량/단순 처리, 비교용 baseline.",
    },
    {
        "name": "tesseract  (LSTM)",
        "combo": "Tesseract 5 LSTM (kor+eng traineddata)",
        "how": "오픈소스 LSTM OCR. 시스템 패키지(sudo) 설치 필요. 레이아웃·표 인식 없음.",
        "pros": ["초경량, 완전 오프라인, 검증된 오픈소스",
                 "설치만 되면 어디서나 동작(CPU)"],
        "cons": ["한국어 문서 정확도 매우 낮음 — CER 33.1% / WER 168.2%(조밀 단락·표 취약)",
                 "레이아웃/표 없음 → 읽기순서 붕괴, 과다 삽입(WER 100% 초과)",
                 "바이너리 sudo 설치 필요(pytesseract + tesseract-ocr-kor)"],
        "fit": "최후의 폴백·대조군, 매우 가벼운 환경.",
    },
]

for e in ENGINES:
    h(doc, e["name"], 2)
    para(doc, "")
    p = doc.add_paragraph()
    p.add_run("구성(모델 조합): ").bold = True
    p.add_run(e["combo"])
    p = doc.add_paragraph()
    p.add_run("동작 방식: ").bold = True
    p.add_run(e["how"])
    para(doc, "장점", bold=True, color=GREEN, size=10)
    bullets(doc, e["pros"])
    para(doc, "단점", bold=True, color=RED, size=10)
    bullets(doc, e["cons"])
    p = doc.add_paragraph()
    p.add_run("적합 상황: ").bold = True
    p.add_run(e["fit"])

# ── 4. 모델별 상세 장표 ──────────────────────────────────────────────────
doc.add_page_break()
h(doc, "4. 모델별 상세 장표 (엔진을 구성하는 개별 모델)", 1)
para(doc, "위 엔진들은 아래 개별 모델들의 조합이다. 각 모델의 출시 시기·특징·장단점.")

MODELS = [
    {
        "name": "Qwen2.5-VL (7B) — 비전언어모델(VLM)",
        "when": "2025년 1월, 알리바바(Alibaba)",
        "sum": "문서 이미지를 통째로 보고 읽는 멀티모달 LLM. 109개 언어, 마크다운/표 출력.",
        "pros": ["자연스러운 줄바꿈·띄어쓰기(가독성·WER 우수)", "강한 한국어 인식",
                 "Ollama로 Blackwell(5090) GPU 우회"],
        "cons": ["비결정적(가끔 전치·누락·환각, 한자 혼입)", "글자 좌표(박스) 없음"],
        "used": "ollama_vl, hybrid_vl(줄글)",
    },
    {
        "name": "PP-StructureV3 — 레이아웃+표+OCR 통합 파이프라인",
        "when": "2025년, PaddleOCR 3.x (Baidu)",
        "sum": "PP-DocLayout + PP-OCRv5 + 표 모델을 묶은 문서 구조 인식 파이프라인.",
        "pros": ["폼·표 구조 인식 최강(셀·도장)", "결정론적, 완전 오프라인", "박스 좌표 제공"],
        "cons": ["한국어 줄글 띄어쓰기 약함(WER↑)", "5090·4090 모두 GPU(cu129, 0.4초/p)"],
        "used": "paddle, hybrid_vl(폼·표)",
    },
    {
        "name": "PP-OCRv5 — 텍스트 검출/인식 모델",
        "when": "2025년, PaddleOCR 3.x",
        "sum": "최신 PP-OCR 세대. 한국어 전용 인식 모델(korean_PP-OCRv5) 제공.",
        "pros": ["빠르고 정확한 글자 인식(CER 우수)", "한국어 전용 모델, 오프라인"],
        "cons": ["줄(line) 단위 인식이라 줄글 띄어쓰기·읽기순서 약함"],
        "used": "paddle, hybrid_vl(폼·표)",
    },
    {
        "name": "PP-DocLayout(_plus-L) — 레이아웃 검출",
        "when": "2024~2025년, PaddleOCR",
        "sum": "문서를 본문/표/그림/각주/머리글/도장 등 23종 영역으로 분류.",
        "pros": ["정밀한 영역 분류(고정밀 plus-L)", "하이브리드 페이지 라우팅의 기반"],
        "cons": ["레이아웃만 — 글자 인식은 별도 모델 필요"],
        "used": "paddle, hybrid_vl(페이지 분류)",
    },
    {
        "name": "SLANeXt / SLANet — 표 구조 복원",
        "when": "2024~2025년, PaddleOCR",
        "sum": "표의 행·열·셀 구조를 복원(유선/무선 표).",
        "pros": ["표 셀 구조 정확 복원", "한국어 셀 텍스트는 메인 OCR과 결합"],
        "cons": ["표가 아닌 영역엔 무관", "복잡·왜곡 표는 한계"],
        "used": "paddle, hybrid_vl(폼·표)",
    },
    {
        "name": "PaddleOCR-VL-0.9B — 경량 비전 OCR",
        "when": "2025년 하반기, PaddleOCR",
        "sum": "0.9B 경량 VLM 단독으로 레이아웃+OCR(서버 불필요, in-process).",
        "pros": ["단독 VLM(외부 서빙 불필요)", "경량·다국어", "5090·4090 GPU 네이티브(cu129)"],
        "cons": ["표-heavy 페이지·헤더 로고 약함", "박스 좌표 없음, 가중치 사전 다운로드 필요"],
        "used": "paddle_vl, hybrid_vl(오프라인 VLM 백엔드)",
    },
    {
        "name": "EasyOCR (CRAFT + CRNN) — 전통 2단계 OCR",
        "when": "2020년경(CRAFT 검출 2019), JaidedAI",
        "sum": "CRAFT로 글자 영역 검출 → CRNN(korean_g2)으로 인식.",
        "pros": ["가볍고 빠름, 설치 간단", "PyTorch 기반 5090 GPU 동작, 오프라인"],
        "cons": ["문서(조밀단락·표) 정확도 낮음(CER 11.4%)", "레이아웃/표 구조 없음"],
        "used": "easyocr",
    },
    {
        "name": "Tesseract 5 (LSTM) — 오픈소스 baseline",
        "when": "2021년 5.0 출시(본 환경 5.3.4). LSTM 인식엔진은 4.0(2018)부터, "
                "원조는 1985년 HP → 2005년 오픈소스화(Google)",
        "sum": "구글 출신 오픈소스 OCR. LSTM 기반, kor+eng traineddata.",
        "pros": ["초경량, 검증된 오픈소스, 완전 오프라인"],
        "cons": ["한국어 문서 정확도 낮음", "레이아웃/표 없음, sudo 설치 필요"],
        "used": "tesseract",
    },
]

# 모델 요약 표
add_table(
    doc,
    ["모델", "출시 시기", "역할", "사용 엔진"],
    [[m["name"].split(" — ")[0], m["when"], m["name"].split(" — ")[-1], m["used"]]
     for m in MODELS],
    center_cols=(),
)
para(doc, "")

for m in MODELS:
    h(doc, m["name"], 2)
    p = doc.add_paragraph()
    p.add_run("출시 시기: ").bold = True
    p.add_run(m["when"])
    p = doc.add_paragraph()
    p.add_run("한 줄 요약: ").bold = True
    p.add_run(m["sum"])
    para(doc, "장점", bold=True, color=GREEN, size=10)
    bullets(doc, m["pros"])
    para(doc, "단점", bold=True, color=RED, size=10)
    bullets(doc, m["cons"])
    p = doc.add_paragraph()
    p.add_run("사용 엔진: ").bold = True
    p.add_run(m["used"])

# ── 5. 평가 방법론 ───────────────────────────────────────────────────────
doc.add_page_break()
h(doc, "5. 평가 방법론 (어떻게 쟀나)", 1)
bullets(doc, [
    ("테스트셋: ", "한국어 계약서 3건 — doc_00(5p)·doc_04(도급계약서 11p)·doc_05(12p), 총 28페이지. 각 페이지 정답(GT) 보유(doc_04·05는 수기 교정)."),
    ("지표: ", "CER(글자 오류율, 공백 무시 — 1차 지표) / WER(단어 오류율, 띄어쓰기 포함 — 가독성 보조)."),
    ("집계: ", "문서 단위(micro) — 5페이지를 이어붙여 CER/WER을 한 번에 계산(웹 멀티엔진 비교와 동일). 페이지별 평균(macro)과 약간 다르며, 본 보고서는 웹과 일치하도록 micro로 통일."),
    ("재현성: ", "paddle·easyocr·tesseract는 완전 결정론(웹과 정확 일치). ollama_vl·hybrid_vl은 seed 고정(OCR_OLLAMA_SEED)으로 ±0.1%p 이내로 안정화 — 완전 bit-동일은 Ollama 잔여 비결정성으로 불가('~' 표기)."),
    ("정규화: ", "곡선/곧은 따옴표, 대시, 원화기호(₩) 등 표기 차이는 동치 처리해 인식 정확도만 평가."),
    ("환경: ", "개발기 RTX 5090(Blackwell) + paddlepaddle-gpu cu129(CUDA 12.9). cu129 휠로 Paddle도 5090에서 GPU(0.4초/p). 배포 4090은 cu126/cu129 둘 다 GPU."),
    ("GT 정리: ", "읽을 수 없는 회사 로고(DONGWOO FINE-CHEM 스타일 이미지)는 GT에서 제외 — 어느 엔진도 못 읽는 비텍스트라 측정 노이즈 제거."),
])
para(doc, "안전장치(병합)란?", bold=True)
para(doc, "줄글 페이지는 항상 'Paddle 글자(정확) + Qwen 띄어쓰기(가독성)'로 문자 단위 병합한다"
          "(`_merge_paddle_chars_qwen_spacing`). 이전엔 큰 누락(hole>=35)일 때만 병합해 Qwen "
          "롤에 따라 CER이 1.5~2.2로 출렁였는데, 항상 병합으로 결정론적 Paddle 글자에 기대 "
          "CER ~1.7%로 안정화(조건부 대비 변동폭 ±0.35 → ±0.1). 전부 Qwen(~2.8%)·전부 Paddle"
          "(WER 43%대) 대비 정확도·가독성 균형. 폼·표 페이지(p0)는 Paddle로 라우팅.")

# ── 6. 최종 결론 ─────────────────────────────────────────────────────────
h(doc, "6. 최종 결론 — 베스트 엔진", 1)
p = doc.add_paragraph()
r = p.add_run("베스트: hybrid_vl (하이브리드 엔진)")
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = GREEN
bullets(doc, [
    ("종합 1위: ", "doc_00·04·05 3개 문서 모두 CER 최저(평균 1.26%) — 정확도 1위 + 가독성 상위(WER)로 종합 최적."),
    ("이유: ", "폼·표는 Paddle(구조·글자 정확), 줄글은 Qwen2.5-VL(자연스러운 띄어쓰기)로 장점만 결합."),
    ("안정성: ", "VLM의 비결정 사고를 자동 감지·병합 복구 → 단일 VLM의 위험을 제거."),
])
para(doc, "상황별 차선책", bold=True)
add_table(
    doc,
    ["상황", "추천 엔진", "근거"],
    [
        ["정확도 최우선(한국어 실문서)", "hybrid_vl", "CER 최저 + 가독성 상위"],
        ["표·폼 위주 + 결정론/오프라인 필수", "paddle", "구조 정확, 사고 없음, 완전 오프라인"],
        ["줄글 위주 + 빠른 GPU 처리", "ollama_vl", "WER 13.6%, 8.4초/p, 5090도 GPU"],
        ["오프라인(Ollama 불가) + VLM 필요", "hybrid_vl + paddle_vl 백엔드", "OCR_HYBRID_VLM=paddle_vl, 평균 CER 2.3%, 서버 불필요"],
        ["5090/4090 단독 VLM", "paddle_vl", "서버 불필요 in-process VLM(GPU)"],
        ["속도·간편 우선(정확도 양보)", "easyocr", "가볍고 빠름, 설치 쉬움"],
    ],
    center_cols=(1,),
)
para(doc, "")
para(doc, "요약: 한국어 문서 OCR의 베스트는 hybrid_vl이다. 표·폼과 줄글이 섞인 실문서에서 "
          "각 부분을 가장 잘하는 모델로 보내고, VLM의 사고는 병합 안전장치로 막아 "
          "정확도와 가독성을 동시에 잡았다.", bold=True)

doc.save(str(OUT))
print(f"[OK] 저장: {OUT}")
print(f"     크기: {OUT.stat().st_size:,} bytes")
